from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "Афонькин М.А. ВКР черновик v8.docx"
OUTPUT = ROOT / "Афонькин М.А. ВКР черновик v9.docx"

STUDENT_FULL = "Афонькин Максим Артемович"
STUDENT_SHORT = "Афонькин М.А."
GROUP = "АВб-22-2"
THEME = "Проектирование и разработка WEB-приложения для рецензирования музыкального творчества"
CITY_YEAR = "Магнитогорск 2026"
CHAIR = "Кафедра вычислительной техники и программирования"
DIRECTION = "Направление 09.03.01 – Информатика и вычислительная техника"
INSTITUTE = "Институт энергетики и автоматизированных систем"
HEAD = "О. С. Логунова"
SUPERVISOR = "Л. Г. Егорова"


def set_run_font(run, size=14, bold=None, italic=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_paragraph(p, *, align=None, first_indent=True, size=14, bold=None, line=1.5, before=0, after=0):
    if align is not None:
        p.alignment = align
    p.paragraph_format.first_line_indent = Cm(1.25) if first_indent else Cm(0)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    for run in p.runs:
        set_run_font(run, size=size, bold=bold)


def add_p(doc, text="", *, align=None, first_indent=False, size=14, bold=None, line=1.0, before=0, after=0):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)
    style_paragraph(
        p,
        align=align,
        first_indent=first_indent,
        size=size,
        bold=bold,
        line=line,
        before=before,
        after=after,
    )
    return p


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    return p


def paragraph_after(doc: Document, anchor, text="", style=None):
    paragraph = doc.add_paragraph()
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    anchor.addnext(paragraph._p)
    return paragraph


def set_doc_defaults(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    for style_name in ("Normal", "Tекст"):
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.size = Pt(14)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            style.paragraph_format.first_line_indent = Cm(1.25)
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.space_after = Pt(0)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.font.bold = True
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.5
    doc.styles["Heading 1"].paragraph_format.first_line_indent = Cm(0)
    doc.styles["Heading 2"].paragraph_format.first_line_indent = Cm(1.25)
    doc.styles["Heading 3"].paragraph_format.first_line_indent = Cm(1.25)


def remove_intro_theme_sentence(doc):
    for p in list(doc.paragraphs):
        if p.text.strip().startswith("Тема выпускной квалификационной работы -"):
            p._element.getparent().remove(p._element)
            break


def remove_existing_front_matter(doc):
    first_body_index = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "СОДЕРЖАНИЕ":
            first_body_index = i
            break
    if first_body_index is None:
        return
    for p in list(doc.paragraphs[:first_body_index]):
        p._element.getparent().remove(p._element)


def normalize_body(doc):
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if p.style.name == "Heading 1":
            if text[:1].isdigit():
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(0)
        elif p.style.name in ("Heading 2", "Heading 3"):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(0)
        elif not text.startswith(("Рисунок ", "Таблица ")):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            set_run_font(run, size=14)
    for table in doc.tables:
        table.alignment = 1
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.first_line_indent = Cm(0)
                    for run in p.runs:
                        set_run_font(run, size=12)


def clear_appendix(doc: Document, heading: str):
    start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == heading:
            start = i
            break
    if start is None:
        return None
    end = len(doc.paragraphs)
    for j in range(start + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].text.strip().startswith("ПРИЛОЖЕНИЕ"):
            end = j
            break
    for p in list(doc.paragraphs[start + 1 : end]):
        p._element.getparent().remove(p._element)
    return doc.paragraphs[start]


def read_code(path: str, max_lines: int | None = None) -> list[str]:
    text = (ROOT.parent / path).read_text(encoding="utf-8")
    lines = text.splitlines()
    if max_lines is not None:
        lines = lines[:max_lines]
    return [line.replace("\t", "    ") for line in lines]


def add_code_listing(doc: Document, anchor_p, title: str, rel_path: str, max_lines: int | None = None):
    title_p = paragraph_after(doc, anchor_p._p, title)
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.first_line_indent = Cm(0)
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after = Pt(3)
    for run in title_p.runs:
        set_run_font(run, size=12, bold=True)
    anchor = title_p
    for line in read_code(rel_path, max_lines=max_lines):
        p = paragraph_after(doc, anchor._p, line if line else " ")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
            run.font.size = Pt(8)
        anchor = p
    return anchor


def rebuild_appendices(doc: Document):
    appendices = [
        (
            "ПРИЛОЖЕНИЕ А СТРУКТУРА API ИНФОРМАЦИОННОЙ СИСТЕМЫ",
            [("Листинг А.1 - Регистрация маршрутов backend API", "backend/routes/routes.go", None)],
        ),
        (
            "ПРИЛОЖЕНИЕ Б ЛИСТИНГ МОДЕЛЕЙ БАЗЫ ДАННЫХ",
            [
                ("Листинг Б.1 - Модель пользователя", "backend/models/user.go", None),
                ("Листинг Б.2 - Модель альбома", "backend/models/album.go", None),
                ("Листинг Б.3 - Модель трека", "backend/models/track.go", None),
                ("Листинг Б.4 - Модель рецензии", "backend/models/review.go", None),
            ],
        ),
        (
            "ПРИЛОЖЕНИЕ В КОНФИГУРАЦИЯ DOCKER COMPOSE",
            [
                ("Листинг В.1 - Конфигурация локального запуска docker-compose.yml", "docker-compose.yml", None),
                ("Листинг В.2 - Конфигурация production-сборки compose.prod.yml", "compose.prod.yml", None),
            ],
        ),
        (
            "ПРИЛОЖЕНИЕ Г КОНФИГУРАЦИЯ GITHUB ACTIONS",
            [("Листинг Г.1 - Pipeline непрерывной интеграции", ".github/workflows/ci.yml", None)],
        ),
    ]
    for heading, listings in appendices:
        anchor = clear_appendix(doc, heading)
        if anchor is None:
            continue
        anchor.style = doc.styles["Heading 1"]
        anchor.alignment = WD_ALIGN_PARAGRAPH.CENTER
        anchor.paragraph_format.first_line_indent = Cm(0)
        for run in anchor.runs:
            set_run_font(run, size=14, bold=True)
        for title, path, max_lines in listings:
            anchor = add_code_listing(doc, anchor, title, path, max_lines=max_lines)


def build_front_matter() -> Document:
    doc = Document()
    set_doc_defaults(doc)

    # Title page.
    add_p(doc, "Министерство науки и высшего образования Российской Федерации", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "Федеральное государственное бюджетное образовательное учреждение", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "высшего образования", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "«Магнитогорский государственный технический университет им. Г.И. Носова»", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_p(doc, "(ФГБОУ ВО «МГТУ им. Г.И. Носова»)", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "")
    add_p(doc, INSTITUTE, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_p(doc, CHAIR, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_p(doc, DIRECTION, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_p(doc, "Допустить к защите", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_p(doc, "Заведующий кафедрой", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_p(doc, f"______________ /{HEAD}/", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_p(doc, "«_____» _______________ 2026 г.", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_p(doc, "", after=18)
    add_p(doc, "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_p(doc, f"обучающегося {STUDENT_FULL}", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"на тему: «{THEME}»", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "", after=12)
    add_p(doc, "ВКР выполнена на ___ страницах", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_p(doc, "Графическая часть на ___ листах", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_p(doc, f"Руководитель: доцент кафедры ВТиП, к.т.н. {SUPERVISOR}", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_p(doc, "(подпись, дата, должность, ученая степень, ученое звание, Ф.И.О)", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_p(doc, "", after=24)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = 1
    table.cell(0, 0).text = "Нормоконтроль и проверка\nна антиплагиат выполнены.\nОригинальность текста ___%\n__________________________ /Л. Г. Егорова/\n(подпись, дата)"
    table.cell(0, 1).text = "Обучающийся________________\n(подпись)\n«_____» _______________ 2026 г."
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if cell is row.cells[1] else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.first_line_indent = Cm(0)
                for run in p.runs:
                    set_run_font(run, size=13)
    add_p(doc, "", after=72)
    add_p(doc, CITY_YEAR, align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)

    # Assignment placeholder.
    add_p(doc, "ЗАДАНИЕ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_p(doc, "на выпускную квалификационную работу", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"Обучающемуся {STUDENT_FULL}", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_p(doc, f"Тема работы: «{THEME}».", align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True, line=1.5)
    add_p(doc, "Срок сдачи обучающимся законченной работы: «_____» _______________ 2026 г.", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_p(doc, "Исходные данные к работе: материалы предметной области рецензирования музыкального творчества, требования к веб-приложению, структура базы данных, программная реализация клиентской и серверной частей системы.", align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True, line=1.5)
    add_p(doc, "Перечень подлежащих разработке вопросов:", align=WD_ALIGN_PARAGRAPH.LEFT)
    for item in [
        "анализ предметной области и существующих веб-платформ рецензирования музыкального творчества;",
        "проектирование архитектуры веб-приложения, модели данных и пользовательских сценариев;",
        "разработка клиентской и серверной частей приложения, административной панели и механизмов модерации;",
        "проверка работоспособности системы, подготовка демонстрационных данных и описание DevOps-инструментов проекта.",
    ]:
        add_p(doc, f"– {item}", align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True, line=1.5)
    add_p(doc, "Задание получил: _________________ /М.А. Афонькин/", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_p(doc, f"Руководитель ВКР: _________________ /{SUPERVISOR}/", align=WD_ALIGN_PARAGRAPH.LEFT)
    page_break(doc)

    # Abstract.
    add_p(doc, "РЕФЕРАТ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_p(doc, f"Тема выпускной квалификационной работы: «{THEME}».", align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True, line=1.5)
    add_p(doc, "Пояснительная записка содержит: ___ страниц, ___ иллюстраций, ___ таблиц, ___ приложений, ___ использованных источников.", align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True, line=1.5)
    add_p(doc, "Ключевые слова: веб-приложение, музыкальное творчество, рецензирование, рейтинговая оценка, модерация, геймификация, профиль пользователя, DevOps, Docker, CI/CD.", align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True, line=1.5)
    abstract = [
        "Актуальность выбранной темы обусловлена потребностью русскоязычной музыкальной аудитории в специализированной платформе, на которой оценки музыкальных произведений сопровождаются аргументированными рецензиями, системой модерации и понятной структурой критериев.",
        "Объектом исследования является процесс коллективного рецензирования музыкального творчества на веб-платформах. Предметом исследования является веб-приложение для публикации рецензий, выставления рейтинговых оценок, работы с музыкальным каталогом и формирования пользовательского профиля.",
        "Целью выпускной квалификационной работы является проектирование и разработка веб-приложения для рецензирования музыкального творчества, обеспечивающего многокритериальную оценку релизов, модерацию пользовательского контента, элементы геймификации и воспроизводимое развертывание программной системы.",
        "В ходе работы выполнен анализ предметной области, спроектирована архитектура информационной системы, разработаны клиентская и серверная части приложения, описана структура базы данных, подготовлены пользовательские сценарии, административные функции и инструменты контейнеризации и автоматизированной проверки проекта.",
    ]
    for text in abstract:
        add_p(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True, line=1.5)
    page_break(doc)

    return doc


def prepend_document(dst: Document, front: Document):
    body = dst._element.body
    first = body[0]
    front_elements = [element for element in front._element.body if element.tag != qn("w:sectPr")]
    for element in front_elements:
        first.addprevious(deepcopy(element))


def main():
    body_doc = Document(SOURCE)
    set_doc_defaults(body_doc)
    remove_existing_front_matter(body_doc)
    remove_intro_theme_sentence(body_doc)
    normalize_body(body_doc)
    rebuild_appendices(body_doc)

    front = build_front_matter()
    prepend_document(body_doc, front)
    set_doc_defaults(body_doc)
    body_doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
