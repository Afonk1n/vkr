from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "Афонькин М.А. ВКР черновик v7.docx"
OUTPUT = ROOT / "Афонькин М.А. ВКР черновик v8.docx"

THEME_OLD = "Разработка информационной системы рецензирования музыкального творчества"
THEME_NEW = "Проектирование и разработка WEB-приложения для рецензирования музыкального творчества"


DB_TABLES = [
    (
        "Пользователи системы (users)",
        [
            ("id", "SERIAL", "Первичный ключ записи пользователя."),
            ("username", "TEXT", "Уникальное имя пользователя, отображаемое в профиле и рецензиях."),
            ("email", "TEXT", "Уникальный адрес электронной почты для входа и идентификации."),
            ("password", "TEXT", "Хеш пароля пользователя; открытое значение пароля не хранится."),
            ("avatar_path", "TEXT", "Путь к изображению аватара пользователя."),
            ("bio", "TEXT", "Краткое описание профиля."),
            ("social_links", "JSONB", "Набор ссылок на социальные сети в структурированном виде."),
            ("is_admin", "BOOLEAN", "Признак наличия административных прав."),
            ("favorite_album_ids", "TEXT", "Список выбранных пользователем любимых альбомов в JSON-представлении."),
            ("favorite_artists", "TEXT", "Список выбранных пользователем любимых артистов в JSON-представлении."),
            ("favorite_track_ids", "TEXT", "Список выбранных пользователем любимых треков в JSON-представлении."),
            ("preferences_manual", "BOOLEAN", "Признак ручного редактирования блока предпочтений."),
            ("is_verified_artist", "BOOLEAN", "Признак верифицированного аккаунта артиста."),
            ("created_at, updated_at, deleted_at", "TIMESTAMPTZ", "Служебные даты создания, изменения и мягкого удаления записи."),
        ],
    ),
    (
        "Жанры музыкальных объектов (genres)",
        [
            ("id", "SERIAL", "Первичный ключ жанра."),
            ("name", "TEXT", "Уникальное название жанра."),
            ("description", "TEXT", "Описание жанра, используемое в каталоге."),
            ("created_at, updated_at, deleted_at", "TIMESTAMPTZ", "Служебные даты жизненного цикла записи."),
        ],
    ),
    (
        "Альбомы (albums)",
        [
            ("id", "SERIAL", "Первичный ключ альбома."),
            ("title", "TEXT", "Название альбома."),
            ("artist", "TEXT", "Имя исполнителя или группы."),
            ("genre_id", "INTEGER", "Внешний ключ на основной жанр альбома."),
            ("cover_image_path", "TEXT", "Путь к обложке альбома."),
            ("release_date", "DATE", "Дата выпуска альбома."),
            ("description", "TEXT", "Краткое описание альбома."),
            ("average_rating", "DOUBLE PRECISION", "Средняя итоговая оценка по одобренным рецензиям."),
            ("created_at, updated_at, deleted_at", "TIMESTAMPTZ", "Служебные даты жизненного цикла записи."),
        ],
    ),
    (
        "Треки (tracks)",
        [
            ("id", "SERIAL", "Первичный ключ трека."),
            ("album_id", "INTEGER", "Внешний ключ на альбом, в который входит трек."),
            ("title", "TEXT", "Название трека."),
            ("duration", "INTEGER", "Продолжительность трека в секундах."),
            ("track_number", "INTEGER", "Порядковый номер трека в альбоме."),
            ("cover_image_path", "TEXT", "Путь к индивидуальной обложке, если она отличается от альбомной."),
            ("average_rating", "DOUBLE PRECISION", "Средняя итоговая оценка трека."),
            ("created_at, updated_at, deleted_at", "TIMESTAMPTZ", "Служебные даты жизненного цикла записи."),
        ],
    ),
    (
        "Связь треков и жанров (track_genres)",
        [
            ("id", "SERIAL", "Первичный ключ связи."),
            ("track_id", "INTEGER", "Внешний ключ на трек."),
            ("genre_id", "INTEGER", "Внешний ключ на жанр."),
            ("UNIQUE(track_id, genre_id)", "Ограничение", "Запрещает повторное назначение одного жанра одному треку."),
        ],
    ),
    (
        "Рецензии и оценки (reviews)",
        [
            ("id", "SERIAL", "Первичный ключ рецензии."),
            ("user_id", "INTEGER", "Внешний ключ на автора рецензии."),
            ("album_id", "INTEGER", "Внешний ключ на альбом; заполняется для рецензии на альбом."),
            ("track_id", "INTEGER", "Внешний ключ на трек; заполняется для рецензии на трек."),
            ("text", "TEXT", "Текстовая часть рецензии."),
            ("rating_rhymes", "INTEGER", "Оценка рифм и образов по шкале от 1 до 10."),
            ("rating_structure", "INTEGER", "Оценка структуры и ритмики по шкале от 1 до 10."),
            ("rating_implementation", "INTEGER", "Оценка реализации стиля по шкале от 1 до 10."),
            ("rating_individuality", "INTEGER", "Оценка индивидуальности и харизмы по шкале от 1 до 10."),
            ("atmosphere_multiplier", "DOUBLE PRECISION", "Множитель атмосферы, влияющий на итоговый балл."),
            ("final_score", "DOUBLE PRECISION", "Итоговая оценка, рассчитанная по формуле системы."),
            ("status", "review_status", "Статус модерации: pending, approved или rejected."),
            ("moderated_by, moderated_at", "INTEGER, TIMESTAMPTZ", "Данные о модераторе и времени проверки."),
            ("created_at, updated_at, deleted_at", "TIMESTAMPTZ", "Служебные даты жизненного цикла записи."),
        ],
    ),
    (
        "Лайки рецензий (review_likes)",
        [
            ("id", "SERIAL", "Первичный ключ реакции."),
            ("user_id", "INTEGER", "Внешний ключ на пользователя, поставившего лайк."),
            ("review_id", "INTEGER", "Внешний ключ на рецензию."),
            ("created_at", "TIMESTAMPTZ", "Дата постановки лайка."),
            ("deleted_at", "TIMESTAMPTZ", "Дата мягкого удаления реакции."),
            ("UNIQUE(user_id, review_id)", "Ограничение", "Не позволяет одному пользователю поставить несколько лайков одной рецензии."),
        ],
    ),
    (
        "Лайки альбомов (album_likes)",
        [
            ("id", "SERIAL", "Первичный ключ реакции."),
            ("user_id", "INTEGER", "Внешний ключ на пользователя."),
            ("album_id", "INTEGER", "Внешний ключ на альбом."),
            ("created_at", "TIMESTAMPTZ", "Дата постановки лайка."),
            ("deleted_at", "TIMESTAMPTZ", "Дата мягкого удаления реакции."),
            ("UNIQUE(user_id, album_id)", "Ограничение", "Не позволяет повторно лайкнуть один альбом одним пользователем."),
        ],
    ),
    (
        "Лайки треков (track_likes)",
        [
            ("id", "SERIAL", "Первичный ключ реакции."),
            ("user_id", "INTEGER", "Внешний ключ на пользователя."),
            ("track_id", "INTEGER", "Внешний ключ на трек."),
            ("created_at", "TIMESTAMPTZ", "Дата постановки лайка."),
            ("deleted_at", "TIMESTAMPTZ", "Дата мягкого удаления реакции."),
            ("UNIQUE(user_id, track_id)", "Ограничение", "Не позволяет повторно лайкнуть один трек одним пользователем."),
        ],
    ),
    (
        "Подписки пользователей (user_follows)",
        [
            ("id", "SERIAL", "Первичный ключ подписки."),
            ("follower_id", "INTEGER", "Внешний ключ на пользователя, который оформил подписку."),
            ("following_id", "INTEGER", "Внешний ключ на пользователя, на которого оформлена подписка."),
            ("created_at", "TIMESTAMPTZ", "Дата создания подписки."),
            ("CHECK(follower_id <> following_id)", "Ограничение", "Запрещает подписку пользователя на самого себя."),
            ("UNIQUE(follower_id, following_id)", "Ограничение", "Запрещает повторную подписку на одного и того же пользователя."),
        ],
    ),
]


def set_font(run, size: int = 14, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def replace_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
    else:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    for run in paragraph.runs:
        set_font(run)


def replace_text_everywhere(doc: Document, old: str, new: str) -> None:
    for paragraph in doc.paragraphs:
        if old in paragraph.text:
            replace_paragraph_text(paragraph, paragraph.text.replace(old, new))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old in paragraph.text:
                        replace_paragraph_text(paragraph, paragraph.text.replace(old, new))


def paragraph_after(doc: Document, element, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    element.addnext(new_p)
    from docx.text.paragraph import Paragraph

    paragraph = Paragraph(new_p, doc._body)
    if style:
        paragraph.style = style
    if text:
        run = paragraph.add_run(text)
        set_font(run)
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph


def table_after(doc: Document, element, rows: int, cols: int):
    table = doc.add_table(rows=rows, cols=cols)
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    element.addnext(tbl)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    return table


def set_cell_border(cell, color: str = "000000", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, margin_twips: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(margin_twips))
        element.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def mark_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = p_pr.find(qn("w:keepNext"))
    if keep_next is None:
        keep_next = OxmlElement("w:keepNext")
        p_pr.append(keep_next)
    keep_next.set(qn("w:val"), "1")


def fill_table(table, headers: list[str], rows: list[tuple[str, str, str]]) -> None:
    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        replace_paragraph_text(cell.paragraphs[0], header)
        for run in cell.paragraphs[0].runs:
            set_font(run, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    mark_header(table.rows[0])

    widths = [3.4, 3.4, 9.0]
    for row_index, row_data in enumerate(rows, start=1):
        for col_index, value in enumerate(row_data):
            cell = table.cell(row_index, col_index)
            replace_paragraph_text(cell.paragraphs[0], value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[col_index])

    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_font(run, size=10)


def insert_physical_tables(doc: Document) -> None:
    caption = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Рисунок 12 -"):
            replace_paragraph_text(paragraph, "Рисунок 12 - Логическая модель базы данных")
            caption = paragraph
        if paragraph.text.strip().startswith("ER-диаграмма отражает"):
            replace_paragraph_text(
                paragraph,
                "Логическая модель отражает основные сущности предметной области и связи между ними: пользователь создает рецензии и реакции, альбом объединяет треки, а рецензия связывает автора с конкретным музыкальным объектом. Такое представление удобно использовать перед переходом к физическому описанию таблиц и внешних ключей.",
            )
            anchor = paragraph._p
            break
    else:
        return

    p = paragraph_after(
        doc,
        anchor,
        "После логической модели сформирована физическая модель базы данных. Она уточняет состав таблиц, типы данных, ключевые поля и назначение каждого атрибута. Такое описание необходимо для проверки соответствия структуры базы данных предметной области и для последующего сопоставления с миграциями PostgreSQL.",
    )
    anchor = p._p
    p = paragraph_after(
        doc,
        anchor,
        "Описание таблиц физической модели приведено в таблицах 4-13.",
    )
    anchor = p._p

    for i, (title, rows) in enumerate(DB_TABLES, start=4):
        page = OxmlElement("w:p")
        run = OxmlElement("w:r")
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run.append(br)
        page.append(run)
        anchor.addnext(page)
        anchor = page
        cap = paragraph_after(doc, anchor, f"Таблица {i} - {title}")
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap.paragraph_format.first_line_indent = Cm(0)
        cap.paragraph_format.space_before = Pt(6)
        cap.paragraph_format.space_after = Pt(6)
        table = table_after(doc, cap._p, rows=len(rows) + 1, cols=3)
        fill_table(table, ["Поле", "Тип данных", "Описание"], rows)
        anchor = table._tbl

    p = paragraph_after(
        doc,
        anchor,
        "Физическая модель показывает, какие данные хранятся постоянно, а какие вычисляются приложением при формировании ответа API. Например, средние оценки по отдельным критериям не сохраняются отдельными столбцами: они рассчитываются по одобренным рецензиям и возвращаются клиентской части вместе с музыкальным объектом.",
    )
    anchor = p._p


def renumber_table_captions_and_refs(doc: Document) -> None:
    caption_re = re.compile(r"^Таблица\s+(\d+)\s+-\s+(.+)$")
    number = 1
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        match = caption_re.match(text)
        if match:
            replace_paragraph_text(paragraph, f"Таблица {number} - {match.group(2)}")
            number += 1

    ref_re = re.compile(r"\b(таблиц(?:е|ы|ах|ами)?|табл\.)\s+(\d+)\b", re.IGNORECASE)

    def repl(match: re.Match) -> str:
        old = int(match.group(2))
        new = old + 10 if old >= 4 else old
        return f"{match.group(1)} {new}"

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if not caption_re.match(text.strip()):
            updated = ref_re.sub(repl, text)
            if updated != text:
                replace_paragraph_text(paragraph, updated)


def center_first_level_headings(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.style.name != "Heading 1":
            continue
        text = paragraph.text.strip()
        if re.match(r"^\d+\s+", text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)


def main() -> None:
    doc = Document(SOURCE)
    replace_text_everywhere(doc, THEME_OLD, THEME_NEW)
    replace_text_everywhere(doc, f"«{THEME_OLD}»", f"«{THEME_NEW}»")
    insert_physical_tables(doc)
    renumber_table_captions_and_refs(doc)
    replace_text_everywhere(doc, "таблицах 14-13", "таблицах 4-13")
    center_first_level_headings(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
