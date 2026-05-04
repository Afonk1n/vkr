from __future__ import annotations

import re
import os
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
CONTENT_PATH = ROOT / "Черновик_ВКР_контент.md"
COURSE_PATH = ROOT / "Афонькин АВб-22-2 Глава 1.docx"
OUTPUT_PATH = Path(os.environ.get("VKR_OUTPUT_DOCX", str(ROOT / "Афонькин М.А. ВКР черновик.docx")))
SCHEMES_ROOT = ROOT / "Схемы" / "png"

ZERO_WIDTH_RE = re.compile("[\u200b\u200c\u200d\u2060\ufeff]")
CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def clean(text: str) -> str:
    text = ZERO_WIDTH_RE.sub("", text)
    text = CONTROL_RE.sub("", text)
    text = text.replace("\xa0", " ")
    text = text.replace("‑", "-").replace("–", "-").replace("—", "-")
    return text.strip()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "BFB3A0", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, margin_twips: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(margin_twips))
        element.set(qn("w:type"), "dxa")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.first_line_indent = Cm(1.25)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr)
    run._r.append(fld_char_2)


def paragraph(doc: Document, text: str = "", style: str | None = None):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(clean(text))
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)
    return p


def title_page(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = paragraph(doc, line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        if line in {"ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", "Магнитогорск, 2026"}:
            p.runs[0].bold = True
        if line.startswith("«") and line.endswith("»"):
            p.runs[0].bold = True


def add_heading(doc: Document, text: str, level: int) -> None:
    has_body = any(p.text.strip() for p in doc.paragraphs)
    if level == 1 and has_body:
        doc.add_page_break()
    p = paragraph(doc, clean(text), style=f"Heading {level}")
    if level == 1:
        p.runs[0].text = clean(text).upper()


def add_caption(doc: Document, text: str) -> None:
    p = paragraph(doc, text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)


def add_table_caption(doc: Document, number: int, text: str) -> None:
    p = paragraph(doc, f"Таблица {number} - {text}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    set_keep_with_next(p)


def add_object_spacer(doc: Document, points: int = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(points)


def set_keep_with_next(p) -> None:
    p_pr = p._p.get_or_add_pPr()
    keep_next = p_pr.find(qn("w:keepNext"))
    if keep_next is None:
        keep_next = OxmlElement("w:keepNext")
        p_pr.append(keep_next)
    keep_next.set(qn("w:val"), "1")


def mark_row_as_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_figure_placeholder(doc: Document, caption: str) -> None:
    image_map = {
        "Архитектура информационной системы": "01_architecture.png",
        "ER-диаграмма базы данных": "02_er_model.png",
        "CI/CD-процесс проекта": "03_cicd.png",
        "Пользовательский сценарий работы с системой": "04_user_flow.png",
        "Сценарий работы администратора": "05_admin_flow.png",
    }
    for marker, filename in image_map.items():
        if marker in caption:
            image_path = SCHEMES_ROOT / filename
            if image_path.exists():
                doc.add_page_break()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(3)
                set_keep_with_next(p)
                run = p.add_run()
                run.add_picture(str(image_path), width=Cm(14.7))
                add_caption(doc, caption)
                return

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_border(cell, "BFB3A0", "8")
    set_cell_shading(cell, "F8F1E7")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    set_keep_with_next(p)
    run = p.add_run("Место для схемы или снимка экрана. Финальное изображение добавляется после утверждения интерфейса.")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.italic = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    add_caption(doc, caption)


def add_simple_table(doc: Document, caption: str, rows: list[list[str]], number: int) -> None:
    add_table_caption(doc, number, caption)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    if rows:
        mark_row_as_table_header(table.rows[0])
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = clean(value)
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(12)
            if r == 0:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
    add_object_spacer(doc)


def copy_course_chapter(doc: Document) -> None:
    course = Document(str(COURSE_PATH))
    started = False
    figure_added_for_next_caption = False
    for src in course.paragraphs:
        text = clean(src.text)
        if not text:
            continue
        if text == "ВВЕДЕНИЕ":
            started = True
        if not started:
            continue
        if text == "БИБЛИОГРАФИЧЕСКИЙ СПИСОК":
            break

        text = text.replace("курсового проекта", "выпускной квалификационной работы")
        text = text.replace("курсовой проект", "выпускная квалификационная работа")
        text = text.replace("КУРСОВОЙ ПРОЕКТ", "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА")
        if text.startswith("Характеристика предметной области"):
            text = "1.1 " + text
        if text == "Выводы":
            text = "1.4 Выводы по первой главе"

        style = src.style.name
        if text.startswith("Рисунок"):
            if not figure_added_for_next_caption:
                add_figure_placeholder(doc, text)
            else:
                add_caption(doc, text)
            figure_added_for_next_caption = False
            continue
        if style.startswith("Heading") or text in {"ВВЕДЕНИЕ"} or re.match(r"^\d+(\.\d+)?\s", text):
            level = 1 if text == "ВВЕДЕНИЕ" or re.match(r"^\d+\s", text) else 2
            add_heading(doc, text, level)
        else:
            p = paragraph(doc, text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def parse_content(doc: Document, content: str) -> None:
    lines = [line.rstrip() for line in content.splitlines()]
    i = 0
    table_counter = 0
    front_matter_done = False
    while i < len(lines):
        line = clean(lines[i])
        if not line:
            i += 1
            continue
        if line == "# TITLE_PAGE":
            i += 1
            while i < len(lines) and not lines[i].startswith("# "):
                i += 1
            continue
        if not front_matter_done:
            if line == "# СОДЕРЖАНИЕ":
                front_matter_done = True
            else:
                i += 1
                continue
        if line == "[COURSE_CHAPTER_1]":
            copy_course_chapter(doc)
            i += 1
            continue
        if line.startswith("[FIGURE:"):
            caption = line.strip("[]").split(":", 1)[1].strip()
            add_figure_placeholder(doc, caption)
            i += 1
            continue
        if line == "[PAGEBREAK]":
            doc.add_page_break()
            i += 1
            continue
        if line.startswith("[TABLE:"):
            caption = line.strip("[]").split(":", 1)[1].strip()
            rows: list[list[str]] = []
            i += 1
            while i < len(lines) and lines[i].startswith("|"):
                parts = [clean(part) for part in lines[i].strip("|").split("|")]
                if not all(set(part) <= {"-", " "} for part in parts):
                    rows.append(parts)
                i += 1
            if rows:
                table_counter += 1
                add_simple_table(doc, caption, rows, table_counter)
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:], 3)
        elif line.startswith("## "):
            add_heading(doc, line[3:], 2)
        elif line.startswith("# "):
            add_heading(doc, line[2:], 1)
        elif line.startswith("- "):
            p = paragraph(doc, line[2:])
            p.style = doc.styles["Normal"]
            p.paragraph_format.left_indent = Cm(0.75)
        else:
            paragraph(doc, line)
        i += 1


def main() -> None:
    doc = Document()
    configure_document(doc)
    add_page_number(doc.sections[0])
    content = CONTENT_PATH.read_text(encoding="utf-8")
    parse_content(doc, content)

    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1)

    doc.save(str(OUTPUT_PATH))


if __name__ == "__main__":
    main()
