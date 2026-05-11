from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

root = Path(r'D:\vkr')
base = next(d for d in root.iterdir() if d.is_dir() and (d / 'CMK-O-SMGTU-36-20.pdf').exists())
out = base / 'db_tables_for_er.docx'
FONT = 'Times New Roman'

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_width(cell, width_cm):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')

def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def style_paragraph(p, size=14, bold=False, align=None, first_line=True):
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    if first_line:
        pf.first_line_indent = Cm(1.25)
    if align is not None:
        p.alignment = align
    for run in p.runs:
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn('w:ascii'), FONT)
        run._element.rPr.rFonts.set(qn('w:hAnsi'), FONT)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        run._element.rPr.rFonts.set(qn('w:cs'), FONT)
        run.font.size = Pt(size)
        run.bold = bold

def add_para(doc, text, size=14, bold=False, align=None, first_line=True):
    p = doc.add_paragraph(text)
    style_paragraph(p, size=size, bold=bold, align=align, first_line=first_line)
    return p

def add_caption(doc, text):
    p = doc.add_paragraph(text)
    style_paragraph(p, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table3(doc):
    rows = [
        ('users', 'Хранение учетных данных, публичного профиля и роли пользователя.', 'Связана с рецензиями, которые создает пользователь.'),
        ('genres', 'Хранение музыкальных жанров, используемых для классификации релизов и треков.', 'Связана с альбомами и треками.'),
        ('albums', 'Хранение музыкальных релизов, обложек, описаний и основного жанра.', 'Связана с жанром, треками и рецензиями.'),
        ('tracks', 'Хранение треков внутри альбома, их порядка, длительности и жанра.', 'Связана с альбомом, жанром и рецензиями.'),
        ('reviews', 'Хранение текста рецензии, статуса модерации и критериев рейтинговой оценки.', 'Связана с пользователем, альбомом или треком.'),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    headers = ['Сущность', 'Назначение', 'Основные связи']
    widths = [3.3, 6.6, 6.4]
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_width(hdr[i], widths[i])
        set_cell_shading(hdr[i], 'F2F2F2')
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(p, size=12, bold=True, first_line=False)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            set_cell_width(cells[i], widths[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                style_paragraph(p, size=11, first_line=False)
    doc.add_paragraph()


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    headers = ['Поле', 'Тип данных', 'Описание']
    widths = [4.1, 3.2, 9.0]
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_width(hdr[i], widths[i])
        set_cell_shading(hdr[i], 'F2F2F2')
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(p, size=12, bold=True, first_line=False)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            set_cell_width(cells[i], widths[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 2 else WD_ALIGN_PARAGRAPH.CENTER
                style_paragraph(p, size=11, first_line=False)
    doc.add_paragraph()

sections = [
    ('Пользователи системы (users)', [
        ('id', 'SERIAL', 'Первичный ключ пользователя.'),
        ('username', 'VARCHAR', 'Уникальное имя пользователя, отображаемое в профиле и рецензиях.'),
        ('email', 'VARCHAR', 'Адрес электронной почты для авторизации и идентификации пользователя.'),
        ('password', 'VARCHAR', 'Хеш пароля пользователя; открытое значение пароля не хранится.'),
        ('avatar_path', 'VARCHAR', 'Путь к изображению аватара пользователя.'),
        ('bio', 'TEXT', 'Краткое описание профиля пользователя.'),
        ('is_admin', 'BOOLEAN', 'Признак наличия административных прав.'),
        ('created_at', 'TIMESTAMP', 'Дата и время регистрации пользователя.'),
    ]),
    ('Жанры музыкальных объектов (genres)', [
        ('id', 'SERIAL', 'Первичный ключ жанра.'),
        ('name', 'VARCHAR', 'Название музыкального жанра.'),
        ('description', 'TEXT', 'Краткое описание жанрового направления.'),
    ]),
    ('Альбомы (albums)', [
        ('id', 'SERIAL', 'Первичный ключ альбома.'),
        ('genre_id', 'INTEGER', 'Внешний ключ на основной жанр альбома.'),
        ('title', 'VARCHAR', 'Название альбома.'),
        ('artist', 'VARCHAR', 'Имя исполнителя или группы.'),
        ('cover_url', 'VARCHAR', 'Путь или URL обложки альбома.'),
        ('description', 'TEXT', 'Краткое описание релиза.'),
        ('created_at', 'TIMESTAMP', 'Дата добавления альбома в каталог.'),
    ]),
    ('Треки (tracks)', [
        ('id', 'SERIAL', 'Первичный ключ трека.'),
        ('album_id', 'INTEGER', 'Внешний ключ на альбом, в который входит трек.'),
        ('genre_id', 'INTEGER', 'Внешний ключ на жанр трека. В упрощенной логической модели используется для прямой связи трека и жанра.'),
        ('title', 'VARCHAR', 'Название трека.'),
        ('duration', 'INTEGER', 'Длительность трека в секундах.'),
        ('track_number', 'INTEGER', 'Порядковый номер трека в альбоме.'),
        ('created_at', 'TIMESTAMP', 'Дата добавления трека.'),
    ]),
    ('Рецензии и оценки (reviews)', [
        ('id', 'SERIAL', 'Первичный ключ рецензии.'),
        ('user_id', 'INTEGER', 'Внешний ключ на автора рецензии.'),
        ('album_id', 'INTEGER', 'Внешний ключ на альбом; заполняется, если рецензия относится к альбому.'),
        ('track_id', 'INTEGER', 'Внешний ключ на трек; заполняется, если рецензия относится к отдельному треку.'),
        ('text', 'TEXT', 'Текстовая часть рецензии.'),
        ('rhyme_score', 'INTEGER', 'Оценка рифм и образов.'),
        ('structure_score', 'INTEGER', 'Оценка структуры и ритмики.'),
        ('style_score', 'INTEGER', 'Оценка реализации стиля.'),
        ('individuality_score', 'INTEGER', 'Оценка индивидуальности и харизмы.'),
        ('vibe_score', 'INTEGER', 'Оценка атмосферы и общего впечатления от музыкального объекта.'),
        ('total_score', 'INTEGER', 'Итоговая рейтинговая оценка, рассчитанная на основе критериев.'),
        ('status', 'VARCHAR', 'Статус модерации рецензии.'),
        ('created_at', 'TIMESTAMP', 'Дата создания рецензии.'),
    ]),
]

doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2)
sec.bottom_margin = Cm(2)
sec.left_margin = Cm(3)
sec.right_margin = Cm(1.5)
styles = doc.styles
styles['Normal'].font.name = FONT
styles['Normal']._element.rPr.rFonts.set(qn('w:ascii'), FONT)
styles['Normal']._element.rPr.rFonts.set(qn('w:hAnsi'), FONT)
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
styles['Normal']._element.rPr.rFonts.set(qn('w:cs'), FONT)
styles['Normal'].font.size = Pt(14)

p = doc.add_paragraph('Фрагмент для раздела 2.3. Физическая модель базы данных')
style_paragraph(p, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
add_para(doc, 'Модель данных включает основные сущности предметной области: пользователей, музыкальные жанры, альбомы, треки и рецензии. Центральной сущностью является рецензия, так как она связывает автора с конкретным музыкальным объектом и хранит результаты многокритериальной оценки.')
add_para(doc, 'Основные сущности базы данных и их связи представлены в таблице 3.')
add_caption(doc, 'Таблица 3 - Основные сущности базы данных')
add_table3(doc)
add_para(doc, 'Физическая структура основных таблиц, соответствующих логической ER-диаграмме, представлена в таблицах 4-8.')
for idx, (caption, rows) in enumerate(sections, start=4):
    if idx in (5, 7, 8):
        doc.add_page_break()
    add_caption(doc, f'Таблица {idx} - {caption}')
    add_table(doc, rows)
doc.save(out)
print(out)
